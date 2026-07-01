from __future__ import annotations

import os
from io import BytesIO
from typing import TYPE_CHECKING, Any

from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

if TYPE_CHECKING:  # apenas para anotações de tipo; não carrega o stack pesado
    from langchain_classic.chains import ConversationalRetrievalChain
    from langchain_core.documents import Document
    from langchain_core.embeddings import Embeddings
    from langchain_huggingface import ChatHuggingFace
    from langchain_mongodb import MongoDBAtlasVectorSearch

load_dotenv()


# ==========================================
# LOADER — leitura de documentos (PDF, DOCX, TXT)
# ==========================================
def carregar_documento(arquivo, nome_arquivo: str) -> list[Document]:
    extensao = nome_arquivo.rsplit(".", 1)[-1].lower()
    if extensao == "pdf":
        return carregar_pdf(arquivo, nome_arquivo)
    if extensao == "docx":
        return carregar_docx(arquivo, nome_arquivo)
    if extensao == "txt":
        return carregar_txt(arquivo, nome_arquivo)
    raise ValueError(f"Extensão não suportada: {extensao}")

def carregar_pdf(arquivo, nome_arquivo: str) -> list[Document]:
    import pdfplumber
    from langchain_core.documents import Document

    documentos = []
    with pdfplumber.open(arquivo) as pdf:
        for numero_pagina, pagina in enumerate(pdf.pages, start=1):
            texto = pagina.extract_text() or ""
            if texto.strip():
                documentos.append(
                    Document(
                        page_content=texto,
                        metadata={"fonte": nome_arquivo, "pagina": numero_pagina},
                    )
                )
    return documentos


def carregar_docx(arquivo, nome_arquivo: str) -> list[Document]:
    import docx
    from langchain_core.documents import Document

    leitor = docx.Document(arquivo)
    texto = "\n".join(p.text for p in leitor.paragraphs if p.text.strip())
    if not texto.strip():
        return []
    return [Document(page_content=texto, metadata={"fonte": nome_arquivo, "pagina": 1})]


def carregar_txt(arquivo, nome_arquivo: str) -> list[Document]:
    from langchain_core.documents import Document

    conteudo = arquivo.read()
    texto = conteudo.decode("utf-8") if isinstance(conteudo, bytes) else conteudo
    if not texto.strip():
        return []
    return [Document(page_content=texto, metadata={"fonte": nome_arquivo, "pagina": 1})]



# ==========================================
# CHUNKER — divisão em chunks
# ==========================================

def dividir_em_chunks(documentos: list[Document], tamanho: int = 500, sobreposicao: int = 50) -> list[Document]:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(chunk_size=tamanho, chunk_overlap=sobreposicao)
    return splitter.split_documents(documentos)


# ==========================================
# EMBEDDER — geração de embeddings
# ==========================================

EMBEDDER_PADRAO = "BAAI/bge-m3"
DIMENSOES_EMBEDDING = 1024  # dimensão do vetor produzido por EMBEDDER_PADRAO


def criar_embedder(modelo: str = EMBEDDER_PADRAO):
    # Embeddings via API da HuggingFace (nuvem) — não exige torch/sentence-transformers
    # localmente. Usa o BAAI/bge-m3 (1024 dims): forte em retrieval multilíngue (PT),
    # sem necessidade de prefixos "query/passage". Trocar o modelo exige reindexar
    # os documentos e recriar o índice de Atlas Search com a nova dimensão.
    from langchain_huggingface import HuggingFaceEndpointEmbeddings

    return HuggingFaceEndpointEmbeddings(
        model=modelo,
        huggingfacehub_api_token=os.environ.get("HUGGINGFACEHUB_API_TOKEN"),
    )


# ==========================================
# VECTOR STORE — indexação e busca (MongoDB Atlas Vector Search)
# ==========================================

NOME_BANCO = "chatbotia"
NOME_COLLECTION = "documentos"
NOME_INDICE_VETORIAL = "vector_index"


def obter_collection():
    from pymongo import MongoClient

    cliente = MongoClient(os.environ["MONGODB_URI"])
    return cliente[NOME_BANCO][NOME_COLLECTION]


def indexar(chunks: list[Document], embedder: Embeddings) -> MongoDBAtlasVectorSearch:
    from langchain_mongodb import MongoDBAtlasVectorSearch

    colecao = obter_collection()
    vector_store = MongoDBAtlasVectorSearch.from_documents(
        documents=chunks,
        embedding=embedder,
        collection=colecao,
        index_name=NOME_INDICE_VETORIAL,
    )
    # from_documents só grava os embeddings; o índice de Atlas Search
    # precisa ser criado explicitamente, senão similarity_search nunca acha nada
    if not list(colecao.list_search_indexes(NOME_INDICE_VETORIAL)):
        vector_store.create_vector_search_index(
            dimensions=DIMENSOES_EMBEDDING, wait_until_complete=60
        )
    return vector_store


def obter_vector_store(embedder: Embeddings | None = None) -> MongoDBAtlasVectorSearch:
    # Conecta ao índice JÁ existente no MongoDB, sem depender de upload na sessão.
    # É o que permite que o chat (admin OU aluno) consulte os documentos que o
    # admin indexou — inclusive depois de reiniciar a API.
    from langchain_mongodb import MongoDBAtlasVectorSearch

    if embedder is None:
        embedder = criar_embedder()
    return MongoDBAtlasVectorSearch(
        collection=obter_collection(),
        embedding=embedder,
        index_name=NOME_INDICE_VETORIAL,
        text_key="text",
        embedding_key="embedding",
    )


def buscar(indice: MongoDBAtlasVectorSearch, pergunta: str, k: int = 4) -> list[Document]:
    return indice.similarity_search(pergunta, k=k)


# ==========================================
# CHAIN — perguntas e respostas (retriever + LLM)
# ==========================================

MODELO_PADRAO = "Qwen/Qwen2.5-7B-Instruct"


def criar_llm(modelo: str = MODELO_PADRAO, temperatura: float = 0.1) -> ChatHuggingFace:
    from langchain_huggingface import ChatHuggingFace, HuggingFaceEndpoint

    endpoint = HuggingFaceEndpoint(
        repo_id=modelo,
        huggingfacehub_api_token=os.environ.get("HUGGINGFACEHUB_API_TOKEN"),
        temperature=temperatura,
        max_new_tokens=512,
    )
    return ChatHuggingFace(llm=endpoint)


def criar_chain(indice: MongoDBAtlasVectorSearch, llm: ChatHuggingFace, k: int = 4) -> ConversationalRetrievalChain:
    from langchain_classic.chains import ConversationalRetrievalChain

    retriever = indice.as_retriever(search_kwargs={"k": k})
    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        return_source_documents=True,
    )


def perguntar(chain: ConversationalRetrievalChain, pergunta: str, historico: list[tuple[str, str]]) -> dict:
    resultado = chain.invoke({"question": pergunta, "chat_history": historico})
    docs = resultado.get("source_documents", [])
    fontes = [
        {"arquivo": doc.metadata.get("fonte"), "pagina": doc.metadata.get("pagina")}
        for doc in docs
    ]
    # Texto dos trechos recuperados — usado pelo revisor para reescrever com fidelidade.
    contexto = "\n\n".join(
        f"[fonte: {doc.metadata.get('fonte')} | pág. {doc.metadata.get('pagina')}]\n{doc.page_content}"
        for doc in docs
    )
    return {"resposta": resultado["answer"], "fontes": fontes, "contexto": contexto}


# ==========================================
# REVISOR — 2ª IA que refina a resposta do RAG
# ==========================================
# Reaproveita o mesmo stack (HuggingFace) — sem nova dependência nem chave.
# Neste projeto o revisor fica SEMPRE ativo (não é desligável por ambiente).
# O modelo do revisor ainda pode ser trocado por env REVISOR_MODELO.

REVISOR_ATIVO = True
# Revisor usa um modelo mais forte que o gerador (7B): o Qwen2.5-72B-Instruct
# entrega reescrita mais fiel e didática. Pode ser trocado por env REVISOR_MODELO.
REVISOR_MODELO = os.environ.get("REVISOR_MODELO", "Qwen/Qwen2.5-72B-Instruct")

REVISOR_SISTEMA = (
    "Você é um tutor de turma. Sua tarefa é reescrever a resposta do sistema para o aluno "
    "de forma clara, didática e acolhedora. Regras:\n"
    "1) Baseie-se SOMENTE no contexto fornecido (os trechos dos documentos). Não invente "
    "fatos que não estejam no contexto; se a informação não constar, diga que não consta "
    "nos documentos.\n"
    "2) Não invente novas fontes. Se citar, use as fontes do contexto.\n"
    "3) Priorize clareza: use frases diretas, passos ou listas quando ajudar.\n"
    "4) Não repita a pergunta nem escreva preâmbulos; responda direto ao aluno.\n"
    "Responda em português do Brasil."
)


def criar_revisor(modelo: str = REVISOR_MODELO, temperatura: float = 0.2):
    from langchain_huggingface import ChatHuggingFace, HuggingFaceEndpoint

    endpoint = HuggingFaceEndpoint(
        repo_id=modelo,
        huggingfacehub_api_token=os.environ.get("HUGGINGFACEHUB_API_TOKEN"),
        temperature=temperatura,
        max_new_tokens=512,
    )
    return ChatHuggingFace(llm=endpoint)


def revisar_resposta(pergunta: str, contexto: str, resposta_bruta: str) -> str:
    from langchain_core.messages import HumanMessage, SystemMessage

    revisor = criar_revisor()
    humano = (
        f"Pergunta do aluno:\n{pergunta}\n\n"
        f"Trechos recuperados (contexto):\n{contexto}\n\n"
        f"Resposta bruta do sistema:\n{resposta_bruta}\n\n"
        "Reescreva a melhor resposta final para o aluno, seguindo as regras."
    )
    saida = revisor.invoke([SystemMessage(content=REVISOR_SISTEMA), HumanMessage(content=humano)])
    texto = getattr(saida, "content", str(saida)).strip()
    return texto or resposta_bruta


# ==========================================
# API — camada HTTP (FastAPI)
# ==========================================

app = FastAPI(title="Chatbot IA API", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class AskRequest(BaseModel):
    question: str = Field(..., min_length=1)
    history: list[tuple[str, str]] | None = None


class AskResponse(BaseModel):
    answer: str
    sources: list[dict[str, Any]]


class UploadResponse(BaseModel):
    message: str
    files: list[str]
    chunks: int


# Estado em memória: guarda o vector store ativo entre requisições.
# Reiniciar o processo zera esta referência (os vetores no MongoDB permanecem,
# mas é preciso reindexar para o /perguntar voltar a respondê-los).
app_state: dict[str, Any] = {
    "arquivos_processados": set(),
    "vector_store": None,
    "total_chunks": 0,
}


@app.get("/saude")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/carregar-documentos", response_model=UploadResponse)
async def upload_documents(files: list[UploadFile] = File(...)) -> UploadResponse:
    if not files:
        raise HTTPException(status_code=400, detail="At least one file is required")

    try:
        processed_files: list[str] = []
        for upload in files:
            if not upload.filename:
                continue

            content = await upload.read()
            stream = BytesIO(content)
            stream.name = upload.filename

            documentos = carregar_documento(stream, upload.filename)
            if not documentos:
                continue

            chunks = dividir_em_chunks(documentos)
            embedder = criar_embedder()

            if app_state["vector_store"] is None:
                app_state["vector_store"] = indexar(chunks, embedder)
            else:
                app_state["vector_store"].add_documents(chunks)

            app_state["arquivos_processados"].add(upload.filename)
            app_state["total_chunks"] += len(chunks)
            processed_files.append(upload.filename)

        if not processed_files:
            raise HTTPException(status_code=400, detail="No valid content found in the uploaded files")

        return UploadResponse(
            message="Documents indexed successfully",
            files=processed_files,
            chunks=app_state["total_chunks"],
        )
    except HTTPException:
        raise
    except KeyError as exc:
        raise HTTPException(status_code=500, detail=f"Missing configuration: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive path
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/documentos")
def list_documents() -> dict[str, Any]:
    # Lista os arquivos efetivamente presentes no vector store (MongoDB),
    # agrupando os chunks por "fonte". É esta a fonte de verdade do que a
    # API consome nas respostas — independe do estado em memória.
    try:
        colecao = obter_collection()
        pipeline = [
            {"$group": {"_id": "$fonte", "chunks": {"$sum": 1}}},
            {"$sort": {"_id": 1}},
        ]
        documentos = [
            {"fonte": grupo["_id"], "chunks": grupo["chunks"]}
            for grupo in colecao.aggregate(pipeline)
            if grupo.get("_id")
        ]
        return {"documentos": documentos}
    except KeyError as exc:
        raise HTTPException(status_code=500, detail=f"Missing configuration: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive path
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.delete("/documentos")
def delete_document(fonte: str) -> dict[str, Any]:
    # Remove do MongoDB todos os chunks de um arquivo (pela metadata "fonte").
    try:
        colecao = obter_collection()
        resultado = colecao.delete_many({"fonte": fonte})
        removidos = resultado.deleted_count
        if removidos == 0:
            raise HTTPException(status_code=404, detail="Documento não encontrado")
        app_state["arquivos_processados"].discard(fonte)
        app_state["total_chunks"] = max(0, app_state["total_chunks"] - removidos)
        return {"fonte": fonte, "removidos": removidos}
    except HTTPException:
        raise
    except KeyError as exc:
        raise HTTPException(status_code=500, detail=f"Missing configuration: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive path
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/estatisticas")
def statistics() -> dict[str, Any]:
    # Estatísticas reais derivadas do vector store (MongoDB) + configuração do RAG.
    try:
        colecao = obter_collection()
        total_chunks = colecao.count_documents({})
        total_documentos = len([f for f in colecao.distinct("fonte") if f])
        return {
            "status": "operational",
            "documentos": total_documentos,
            "chunks": total_chunks,
            "modelo_saida": MODELO_PADRAO,
            "modelo_resposta": REVISOR_MODELO,
            "embedder": EMBEDDER_PADRAO,
            "banco_vetorial": "MongoDB Atlas",
            "dimensao": DIMENSOES_EMBEDDING,
            "estrategia": "Vector Search",
            "indice": NOME_INDICE_VETORIAL,
        }
    except KeyError as exc:
        raise HTTPException(status_code=500, detail=f"Missing configuration: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive path
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/perguntar", response_model=AskResponse)
def ask_question(payload: AskRequest) -> AskResponse:
    # Consulta SEMPRE o índice persistido no MongoDB (não o estado em memória).
    # Assim, tanto o chat do admin quanto o do aluno respondem com base nos
    # documentos que o admin indexou — mesmo após reiniciar a API.
    try:
        if obter_collection().count_documents({}) == 0:
            raise HTTPException(status_code=400, detail="No documents have been indexed yet")

        vector_store = obter_vector_store()
        llm = criar_llm()
        chain = criar_chain(vector_store, llm)
        resultado = perguntar(chain, payload.question, payload.history or [])

        resposta_final = resultado["resposta"]
        # 2ª IA (revisor): refina a resposta com base no contexto recuperado.
        # Se falhar (cota/timeout/erro), cai de volta para a resposta bruta.
        if REVISOR_ATIVO and resultado.get("contexto"):
            try:
                resposta_final = revisar_resposta(
                    payload.question, resultado["contexto"], resultado["resposta"]
                )
            except Exception:  # pragma: no cover - fallback defensivo
                resposta_final = resultado["resposta"]

        return AskResponse(answer=resposta_final, sources=resultado["fontes"])
    except HTTPException:
        raise
    except KeyError as exc:
        raise HTTPException(status_code=500, detail=f"Missing configuration: {exc}") from exc
    except Exception as exc:  # pragma: no cover - defensive path
        raise HTTPException(status_code=500, detail=str(exc)) from exc
